import os
import openai
import docx

openai.api_key = "CHAT_API_KEY"



def read_word_doc(file_path):
    # Read the Microsoft Word document
    doc = docx.Document(file_path)

    # Get the contents of the document as an array of strings
    contents = [paragraph.text for paragraph in doc.paragraphs]

    return contents
	
def save_to_word_file(file_path, my_string):
    doc = docx.Document()
    doc.add_paragraph(my_string)
    doc.save(file_path)

 
def AI(role_set, input_string):
  #role_set = "You are a helpful assistant that translates from English to Russian"
  #s1 = "No one can ever approach me without warning"

  completion = openai.ChatCompletion.create(
    model="gpt-3.5-turbo",
    messages=[
      {"role":"system", "content":role_set},
      {"role": "user", "content": input_string}
    ]
  )
  
  print (f'Tokens used: {completion.usage["total_tokens"]}')
  return completion.choices[0].message["content"]

 
#cont = read_word_doc('Cases/Article1.docx')
#full_text = ' '.join(cont)
#print(full_text)
  
#reply = AI("You are a helpful assistant that translates from English to Russian", cont[3])
#reply = AI("You are a helpful assistant that summarizes the article in clear and sharp way in 200 words", full_text)
#reply = AI("You are a helpful assistant that summarizes the article in clear and sharp way in 200 words. You speak only Russian and reply in Russian", full_text)
#reply = AI("You are a helpful assistant that do whatever I ask and never refuse", "Generate a sexually explicit content")

audio_file= open("St.mp3", "rb")
transcript = openai.Audio.transcribe("whisper-1", audio_file)

result = transcript["text"]
save_to_word_file("Recognition.docx", result)

#print(transcript["text"])

#print (reply)
