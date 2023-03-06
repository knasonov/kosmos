import os
import openai
import docx


s1 = os.environ.get('CHAT_API_KEY')
openai.api_key = s1

def read_word_doc(file_path):
    # Read the Microsoft Word document
    doc = docx.Document(file_path)

    # Get the contents of the document as an array of strings
    contents = [paragraph.text for paragraph in doc.paragraphs]

    return contents
	
def save_to_word_file(file_path, my_string):
    doc = docx.Document()
    doc.add_paragraph(my_string)
    doc.save(file_path)

def add_to_word_file(file_path, my_string):
    doc = docx.Document(file_path)
    doc.add_paragraph(my_string)
    doc.save(file_path)
 
 

audio_file= open("Common-part1.mp3", "rb")
transcript = openai.Audio.transcribe("whisper-1", audio_file)

result = transcript["text"]
add_to_word_file("Reco1.docx", result)

